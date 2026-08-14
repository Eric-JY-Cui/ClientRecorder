import os
from docx import Document
from docx2pdf import convert
import datetime
from num2words import num2words

def generate_receipt(information:dict, file_name:str):
    """
    generate receipt in both docx and pdf version based on file input. 
    """
    template_path = "reciept_template.docx"
    output_dir = "Reciepts"
    # Ensure output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    new_filename_base = file_name
        
    # Define output file paths
    docx_output_path = os.path.join(output_dir, information["name"],f"{new_filename_base}.docx")
    pdf_output_path = os.path.join(output_dir, information["name"], f"{new_filename_base}.pdf")
    os.makedirs(os.path.join(output_dir, information["name"]), exist_ok=True)

    template_dict = {
        "<name>" : information["name"],
        "<address>" : information["address"],
        "<health_issue>" : information["health_issue"],
        "<total_cost>" : information["cost"],
        "<total_printed>" : num2words(int(information["cost"])).title() + " Dollars.",
        "<date>" : information["date"]
    }

    for i in range(10):
        if i < len(information["visits"]):
            template_dict[f"<v{i}year>"] = information["visits"][i]["year"]
            template_dict[f"<v{i}month>"] = information["visits"][i]["month"]
            template_dict[f"<v{i}day>"] = information["visits"][i]["day"]
            template_dict[f"<v{i}TMmessage>"] = information["visits"][i]["modality"]
            template_dict[f"<v{i}duration>"] = str(information["visits"][i]["duration"]) + " mins"
            template_dict[f"<v{i}cost>"] = "$" + str(information["visits"][i]["cost"])
        else:
            template_dict[f"<v{i}year>"] = ""
            template_dict[f"<v{i}month>"] = ""
            template_dict[f"<v{i}day>"] = ""
            template_dict[f"<v{i}TMmessage>"] = ""
            template_dict[f"<v{i}duration>"] = ""
            template_dict[f"<v{i}cost>"] = ""

    print(template_dict)
    
    try:
        # 1. Open the template document
        doc = Document(template_path)
        for key,value in template_dict.items():
            find_text = key
            replace_text = value
            # 2. Search and replace in paragraphs (main body text)
            for paragraph in doc.paragraphs:
                if find_text in paragraph.text:
                    for run in paragraph.runs:
                        if find_text in run.text:
                            run.text = run.text.replace(find_text, str(replace_text))
                            
            # 3. Search and replace in tables (if your keyword is inside a table)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if find_text in paragraph.text:
                                for run in paragraph.runs:
                                    if find_text in run.text:
                                        run.text = run.text.replace(find_text, str(replace_text))

        
        # 4. Save the modified document as a new Word file
        doc.save(docx_output_path)
        print(f"Successfully saved Word doc: {docx_output_path}")
        
        # 5. Convert the new Word file to PDF
        print("Converting to PDF (this may take a moment)...")
        convert(docx_output_path, pdf_output_path)
        print(f"Successfully saved PDF: {pdf_output_path}")
        
    except Exception as e:
        print(f"An error occurred: {e}")


if __name__ == "__main__":
    info = {'name': 'Robert Chen', 'address': '42 Maple Road, Laval, QC H7N 3C2', 'health_issue': 'testing', 'cost': 400.0, 'date': datetime.date(2026, 8, 13), 'visits': [{'year': 2026, 'month': 8, 'day': 8, 'modality': 'Ultrasound', 'duration': 30, 'cost': 125}, {'year': 2026, 'month': 6, 'day': 12, 'modality': 'Manual Therapy', 'duration': 75, 'cost': 200}, {'year': 2026, 'month': 5, 'day': 13, 'modality': 'Therapeutic Exercise', 'duration': 45, 'cost': 75}]}
    generate_receipt(info,"test")